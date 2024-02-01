from platform import python_compiler
from flask import Flask, render_template, request, send_file
from PIL import Image
from io import BytesIO
from docx import Document
from docx2pdf import convert
from flask_cors import CORS
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from wand.image import Image

app = Flask(__name__)
CORS(app, origins="http://localhost:5173")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert/png2jpg', methods=['POST'])
def png_to_jpg():
    return convert_image('JPEG', 'converted_image.jpg')

@app.route('/convert/jpg2png', methods=['POST'])
def jpg_to_png():
    return convert_image('PNG', 'converted_image.png')

@app.route('/convert/png2jpeg', methods=['POST'])
def png_to_jpeg():
    return convert_image('JPEG', 'converted_image.jpg')

@app.route('/convert/jpeg2png', methods=['POST'])
def jpeg_to_png():
    return convert_image('PNG', 'converted_image.png')

def convert_image(format, download_name):
    uploaded_file = request.files['file']

    if uploaded_file.filename != '':
        image = Image.open(uploaded_file)

        if image.mode == 'RGBA':
            image = image.convert('RGB')

        buffer = BytesIO()

        # Save the image in 'JPEG' format if the target format is 'JPEG', otherwise 'PNG'
        image.save(buffer, format='JPEG' if format == 'JPEG' else 'PNG')

        buffer.seek(0)
        
        # Set the file extension based on the target format
        if format == 'JPEG':
            download_name += '.jpg'
        else:
            download_name += '.png'

        return send_file(buffer, download_name=download_name, as_attachment=True, mimetype=f'image/{format.lower()}')

    return "No file selected"


def convert_docx_to_pdf(docx_content):
    pdf_buffer = BytesIO()

    doc = Document(BytesIO(docx_content))
    
    # Create a ReportLab PDF document
    pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    
    # Define custom styles matching Word document
    custom_styles = getSampleStyleSheet()
    custom_paragraph_style = ParagraphStyle(
        'CustomParagraph',
        parent=custom_styles['Normal'],
        fontSize=12,  # Adjust this value to match the font size in your Word document
        spaceAfter=12,  # Adjust this value to match the spacing between paragraphs
    )
    
    # Create flowables for paragraphs using custom styles
    flowables = []
    for para in doc.paragraphs:
        flowables.append(Paragraph(para.text, custom_paragraph_style))
    
    # Build the PDF document
    pdf_doc.build(flowables)
    
    pdf_buffer.seek(0)
    return pdf_buffer


@app.route("/convert/word2pdf", methods=["POST"])
# app.route("/convert/word-to-pdf", methods=["POST"])
def convert():
    if 'file' not in request.files:
        return 'No file part'

    file = request.files['file']

    if file.filename == '':
        return 'No selected file'

    if file and file.filename.endswith('.docx'):
        docx_content = file.read()
        pdf_buffer = convert_docx_to_pdf(docx_content)
        print("success")
        return send_file(pdf_buffer, download_name='output.pdf', as_attachment=True)

    return 'Invalid file format. Please upload a .docx file.'

def convert_docx_to_jpg(docx_content):
    # Create a ReportLab PDF document
    pdf_buffer = BytesIO()
    doc = Document(BytesIO(docx_content))
    
    pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)

    # Corrected import for getSampleStyleSheet and ParagraphStyle
    custom_styles = getSampleStyleSheet()
    custom_paragraph_style = ParagraphStyle(
        'CustomParagraph',
        parent=custom_styles['Normal'],
        fontSize=8,  # Adjust this value for smaller font size
        spaceAfter=12,
    )

    # Create flowables for paragraphs using custom styles
    flowables = []
    for para in doc.paragraphs:
        flowables.append(Paragraph(para.text, custom_paragraph_style))

    pdf_doc.build(flowables)
    pdf_buffer.seek(0)

    # Convert the PDF to an image buffer (JPEG) with improved quality settings
    with Image(blob=pdf_buffer.read(), format='pdf', resolution=300) as pdf_image:
        # Set density to control the rendering quality
        pdf_image.options.density = 8000  # Adjust this value for better clarity

        jpg_buffer = BytesIO(pdf_image.make_blob('jpeg'))

    pdf_buffer.close()
    jpg_buffer.seek(0)
    return jpg_buffer


@app.route('/convert/word2jpg', methods=['POST'])
def convert_to_jpg():
    if 'file' not in request.files:
        return 'No file part'

    file = request.files['file']

    if file.filename == '':
        return 'No selected file'

    if file and file.filename.endswith('.docx'):
        docx_content = file.read()
        
        # Convert to JPG
        jpg_buffer = convert_docx_to_jpg(docx_content)
        return send_file(jpg_buffer, download_name='output.jpg', as_attachment=True)

    return 'Invalid file format. Please upload a .docx file.'


if __name__ == '__main__':
    app.run(debug=True)
